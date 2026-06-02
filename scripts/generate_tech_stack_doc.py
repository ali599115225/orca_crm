# -*- coding: utf-8 -*-
"""Generate ORCA CRM technology stack Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

OUT = r"C:\Users\ali59\Desktop\REDC\docs\ORCA_CRM_التقنيات_المستخدمة.docx"

sections = [
    ("1. نظرة عامة على المنصة", [
        "ORCA CRM — منصة SaaS سحابية لإدارة العمليات العقارية (مبيعات، مشاريع، إيجارات، تسويق، وكلاء ذكاء اصطناعي).",
        "النطاق الإنتاجي: orca.az-ez.pro | البنية: تطبيق ويب Full-Stack على Next.js.",
    ]),
    ("2. الواجهة الأمامية (Frontend)", [
        ("Next.js", "إطار React مع App Router لتوجيه الصفحات وعرض الخادم SSR/SSG"),
        ("React 19", "مكتبة بناء واجهات المستخدم التفاعلية"),
        ("TypeScript", "لغة البرمجة مع التحقق الثابت من الأنواع"),
        ("Tailwind CSS v4", "إطار تنسيق CSS utility-first"),
        ("PostCSS + Autoprefixer", "معالجة CSS ودعم المتصفحات"),
        ("Phosphor Icons", "مكتبة أيقونات عبر CDN (@phosphor-icons/web)"),
        ("Lucide React", "أيقونات إضافية في بعض المكوّنات"),
        ("GSAP", "مكتبة حركات وانتقالات متقدمة (صفحات تسويقية)"),
        ("Google Fonts", "خط Cairo للعربية و Inter للإنجليزية والأرقام"),
        ("دعم RTL/LTR", "واجهة ثنائية اللغة (عربي / إنجليزي) مع تبديل الاتجاه"),
        ("الوضع الداكن", "Dark mode مدمج عبر React Context"),
    ]),
    ("3. الخادم والتطبيق (Backend)", [
        ("Next.js App Router", "صفحات Server Components + Client Components"),
        ("Server Actions", "منطق الخادم (app/actions/*) دون API منفصل لمعظم العمليات"),
        ("REST API Routes", "app/api/* — نقاط نهاية REST للتكامل الخارجي"),
        ("Node.js 20+", "بيئة تشغيل الخادم"),
        ("Proxy (proxy.ts)", "طبقة وسيطة للطلبات والحماية"),
    ]),
    ("4. قاعدة البيانات وطبقة البيانات", [
        ("PostgreSQL", "قاعدة بيانات علائقية رئيسية"),
        ("Prisma ORM v7", "نمذجة البيانات، migrations، واستعلامات type-safe"),
        ("@prisma/adapter-pg", "محول PostgreSQL لـ Prisma 7"),
        ("@prisma/adapter-neon", "دعم Neon Serverless"),
        ("@neondatabase/serverless", "اتصال serverless بقاعدة Neon"),
        ("pg (node-postgres)", "عميل PostgreSQL مباشر للـ connection pool"),
        ("Multi-Tenant", "عزل بيانات المستأجرين عبر tenant_id + Prisma Extension"),
        ("Audit Logs", "سجل تدقيق تلقائي للعمليات على قاعدة البيانات"),
    ]),
    ("5. المصادقة والأمان", [
        ("JWT (jose)", "جلسات المستخدم وتوقيع الرموز"),
        ("bcryptjs", "تشفير كلمات المرور (Salt Rounds)"),
        ("AES-256-CBC (crypto)", "تشفير البيانات الحساسة at-rest (مفاتيح API، ZATCA)"),
        ("عزل المستأجرين", "tenant-context + فرض tenantId على الاستعلامات"),
        ("أدوار RBAC", "ADMIN, SALES_MANAGER, SALES_EMPLOYEE, MARKETING, READ_ONLY, PLATFORM_ARCHITECT"),
        ("HTTPS / SSL", "اتصال مشفّر مع قاعدة البيانات"),
    ]),
    ("6. الاستضافة والنشر (DevOps)", [
        ("Vercel", "استضافة serverless ونشر تلقائي من Git"),
        ("Vercel Cron Jobs", "مهام مجدولة: الفوترة (02:00) و Sentinel (06:00)"),
        ("Git / GitHub", "إدارة الإصدارات (orca_crm)"),
        ("متغيرات البيئة", "DATABASE_URL, JWT_SECRET, DIRECT_URL, مفاتيح التكامل"),
        ("Node scripts", "رفع env إلى Vercel، seed، اختبارات scratch"),
    ]),
    ("7. التكاملات الخارجية", [
        ("Moyasar", "بوابة دفع سعودية (اشتراكات SaaS — فواتير بالريال)"),
        ("Resend", "إرسال البريد الإلكتروني التنبيهي للمسؤولين"),
        ("WhatsApp Webhook", "app/api/whatsapp/webhook — قناة واتساب للعملاء"),
        ("Leads Webhook API", "app/api/v1/leads/webhook — استقبال عملاء من الحملات"),
        ("ZATCA", "امتثال الفاتورة الإلكترونية (المرحلة الثانية) — واجهات وحفظ مشفّر"),
        ("Ejar", "تكامل عقود الإيجار الحكومية (وحدة ejar في النظام)"),
        ("منصات إعلانية", "Google, Meta, TikTok, Snapchat, X, LinkedIn — ربط حسابات الإعلان"),
    ]),
    ("8. الوكلاء الذكيون (AI Agents)", [
        ("ساهر (Saher)", "وكيل الامتثال — فحص عقود وفواتير"),
        ("سند (Sanad)", "وكيل الفوترة والاشتراكات السحابية"),
        ("بصير (Baseer)", "تحليل ROI والتسويق الاستباقي"),
        ("منصور (Mansour)", "أتمتة واتساب ومسارات المتابعة"),
        ("خبير (Khabeer)", "دعم وأتمتة قانونية"),
        ("Agent Slots & Leases", "إدارة تراخيص الوكلاء وعداد الاستخدام"),
        ("Telemetry Logs", "سجلات تشغيل الوكلاء"),
        ("NLP / Lead Scoring", "تقييم جدية العملاء من مصدر الإعلان"),
    ]),
    ("9. وحدات الأعمال في المنصة", [
        "لوحة التحليلات والمؤشرات (Dashboard KPIs)",
        "إدارة العملاء المحتملين (Leads) — Kanban و Table",
        "إدارة المشاريع العقارية والوحدات (Projects / Units)",
        "الإيجارات والمحاسبة ERP (عقود، أقساط، ZATCA)",
        "حاسبة التمويل السكني",
        "أداء المبيعات والعمولات",
        "لوحة النمو والتسويق (ROI، قنوات، واتساب)",
        "إدارة الوكلاء والذكاء الاصطناعي",
        "المهام والتذكيرات",
        "مركز الدعم (Helpdesk + تذاكر)",
        "قناة واتساب",
        "سجل النظام (Logs)",
        "إعدادات المنشأة والمستخدمين",
        "عقود الحجز الإلكترونية",
        "تسجيل المنشآت والاشتراكات (SaaS Onboarding)",
    ]),
    ("10. واجهات API الرئيسية", [
        "GET/POST /api/v1/auth/login — مصادقة",
        "GET /api/v1/health — فحص صحة النظام",
        "GET /api/v1/dashboard/metrics — مؤشرات لوحة التحكم",
        "GET /api/v1/dashboard/units — مخزون الوحدات العقارية",
        "GET /api/v1/dashboard/telemetry — سجلات الوكلاء",
        "POST /api/v1/leads/webhook — استقبال leads خارجي",
        "POST /api/whatsapp/webhook — واتساب",
        "GET/POST /api/projects, /api/leads, /api/tasks",
        "GET /api/payment/callback — رد بوابة الدفع",
        "GET /api/cron/billing, /api/cron/sentinel, /api/cron/installments",
    ]),
    ("11. أدوات التطوير", [
        ("Prisma Migrate", "ترحيل مخطط قاعدة البيانات"),
        ("Prisma Seed", "بيانات تجريبية أولية"),
        ("tsx", "تشغيل سكربتات TypeScript"),
        ("dotenv", "إدارة متغيرات البيئة محلياً"),
        ("TypeScript + @types/*", "تعريفات الأنواع"),
        ("ESLint / build", "prisma generate && next build"),
    ]),
]

def add_table_section(doc, title, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "التقنية / المكوّن"
    hdr[1].text = "الوصف والاستخدام"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        if isinstance(row, tuple):
            cells[0].text = row[0]
            cells[1].text = row[1]
        else:
            cells[0].merge(cells[1])
            cells[0].text = row
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("تقرير التقنيات المستخدمة", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("منصة ORCA CRM — نظام إدارة العمليات العقارية السحابي")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    date_p = doc.add_paragraph(f"تاريخ الإعداد: {date.today().strftime('%Y-%m-%d')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for title_text, content in sections:
        if content and isinstance(content[0], tuple):
            add_table_section(doc, title_text, content)
        else:
            doc.add_heading(title_text, level=1)
            for item in content:
                p = doc.add_paragraph(item, style="List Bullet")
                for r in p.runs:
                    r.font.size = Pt(11)

    doc.add_page_break()
    doc.add_heading("ملاحظة", level=1)
    doc.add_paragraph(
        "هذا المستند يعكس حالة المشروع البرمجي (REDC / ORCA CRM) كما هي في المستودع الحالي. "
        "قد تُضاف أو تُحدَّث بعض التكاملات حسب إعدادات متغيرات البيئة في الإنتاج."
    )

    doc.save(OUT)
    print(f"Created: {OUT}")

if __name__ == "__main__":
    main()
